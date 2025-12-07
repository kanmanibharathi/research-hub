# File: D:\GitHub\stat-summary-webapp\build_doc.py
import os
import pandas as pd
from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(__file__)
SUMMARY_CSV = os.path.join(BASE, "summary.csv")
PLOTS_DIR = os.path.join(BASE, "plots")
OUT_DOC = os.path.join(BASE, "Summary_A4_landscape.docx")

def set_a4_landscape(section):
    # A4 size in inches: 8.27 x 11.69 -> landscape width=11.69, height=8.27
    section.page_width = Mm(297)   # 297 mm = 11.6929 in
    section.page_height = Mm(210)  # 210 mm = 8.2677 in
    # set small margins to maximize content
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

def insert_table(document, df):
    # Create table with header row
    nrows, ncols = df.shape[0] + 1, df.shape[1]
    table = document.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    # header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        para = hdr_cells[i].paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(8)
    # data rows
    for r_idx, row in df.iterrows():
        row_cells = table.rows[r_idx+1].cells
        for c_idx, col in enumerate(df.columns):
            row_cells[c_idx].text = str(row[col])
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if isinstance(row[col], (int,float)) else WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in p.runs:
                run.font.size = Pt(8)
    # tighten up table width
    table.autofit = True

def insert_plots_row(document, df):
    # Insert a single row of images to the right side: we will place a caption row with small images next to the table
    # Simpler approach: add a paragraph and insert each image inline with width set.
    imgs = []
    for var in df['Variable']:
        png = os.path.join(PLOTS_DIR, f"{var}.png")
        if os.path.exists(png):
            imgs.append(png)
        else:
            imgs.append(None)
    # Put images in a narrow table (two columns) next to table is more complex. We'll append at end and set each image small.
    document.add_page_break()  # remove if we want everything on one page; but we'll instead place images in a 1-row table
    # Alternative: add a 1-column table with image per row after the summary table
    for i, img in enumerate(imgs):
        if img:
            p = document.add_paragraph()
            run = p.add_run()
            # width: choose a width so that combined table + images fit A4 landscape:
            run.add_picture(img, width=Inches(1.0))  # small sparkline-sized images
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def build():
    if not os.path.exists(SUMMARY_CSV):
        raise FileNotFoundError("Run generate_summary.py first to create summary.csv")
    df = pd.read_csv(SUMMARY_CSV)
    # reorder columns to a nicer readable order if needed
    doc = Document()
    # set A4 landscape for first section
    section = doc.sections[0]
    set_a4_landscape(section)
    # Title
    title = doc.add_paragraph("Summary Statistics")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.size = Pt(14)
    title.runs[0].font.bold = True
    # Insert table
    insert_table(doc, df)
    # After table, add a small spacer and then add the histograms (one small per variable)
    doc.add_paragraph()
    # Insert histograms in a single row table: create table with 1 row and n columns
    n = len(df)
    if n > 0:
        tbl = doc.add_table(rows=1, cols=n)
        tbl.autofit = False
        # compute image width to fit page width:
        page_width = section.page_width - section.left_margin - section.right_margin
        # pick image width ~ page_width/n (but not exceed 2.0 inches)
        max_inch = Mm(25)  # roughly 1 inch
        # convert Mm to inches: python-docx accepts Inches, but we'll compute Inches(len)
        for i, var in enumerate(df['Variable']):
            cell = tbl.rows[0].cells[i]
            cell.width = Mm( (page_width / n).mm if hasattr(page_width,'mm') else 30 )
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pic_path = os.path.join(PLOTS_DIR, f"{var}.png")
            if os.path.exists(pic_path):
                run = p.add_run()
                run.add_picture(pic_path, width=Inches(1.0))
    # Save doc
    doc.save(OUT_DOC)
    print("Saved DOCX:", OUT_DOC)

if __name__ == "__main__":
    build()
