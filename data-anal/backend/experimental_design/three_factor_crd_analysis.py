import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from .duncan_util import get_duncan_q

class ThreeFactorCRDAnalyzer:
    def __init__(self, df, a_col, b_col, c_col, resp_col):
        self.df = df
        self.a_col = a_col
        self.b_col = b_col
        self.c_col = c_col
        self.resp_col = resp_col
        
        self.a = 0
        self.b = 0
        self.c = 0
        self.n = 0
        
        self.anova_table = {}
        self.results = {}
        
        self.MS_E = 0
        self.df_E = 0
        self.grand_mean = 0
        self.r_bar = 0
        self.alpha = 0.05

    def validate(self):
        # Type conversion
        self.df[self.a_col] = self.df[self.a_col].astype(str)
        self.df[self.b_col] = self.df[self.b_col].astype(str)
        self.df[self.c_col] = self.df[self.c_col].astype(str)
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        
        if self.df[self.resp_col].isnull().any():
             raise ValueError("Response variable contains missing or non-numeric values.")
             
        self.n = len(self.df)
        self.a = self.df[self.a_col].nunique()
        self.b = self.df[self.b_col].nunique()
        self.c = self.df[self.c_col].nunique()
        
        # Check DF
        if self.n <= self.a * self.b * self.c:
            raise ValueError(f"Insufficient degrees of freedom. Observations ({self.n}) must exceed treatments ({self.a * self.b * self.c}).")

    def run_anova(self):
        y = self.df[self.resp_col]
        G = y.sum()
        self.grand_mean = y.mean()
        CT = (G**2) / self.n
        USS_Total = (y**2).sum()
        SS_Total = USS_Total - CT
        
        # Determine group sums and counts (for robust unbalanced handling if needed, though assume balanced often)
        # Using helper for USS calc
        def get_uss(cols):
            grp = self.df.groupby(cols)[self.resp_col]
            # sum( sum(y)^2 / n_cell )
            return sum((g.sum()**2 / len(g)) for _, g in grp)

        USS_A = get_uss([self.a_col])
        USS_B = get_uss([self.b_col])
        USS_C = get_uss([self.c_col])
        
        USS_AB = get_uss([self.a_col, self.b_col])
        USS_AC = get_uss([self.a_col, self.c_col])
        USS_BC = get_uss([self.b_col, self.c_col])
        
        USS_ABC = get_uss([self.a_col, self.b_col, self.c_col])
        
        # SS Calculations
        SS_A = USS_A - CT
        SS_B = USS_B - CT
        SS_C = USS_C - CT
        
        SS_AB = USS_AB - CT - SS_A - SS_B
        SS_AC = USS_AC - CT - SS_A - SS_C
        SS_BC = USS_BC - CT - SS_B - SS_C
        
        SS_ABC = USS_ABC - CT - SS_A - SS_B - SS_C - SS_AB - SS_AC - SS_BC
        
        SS_Error = SS_Total - (SS_A + SS_B + SS_C + SS_AB + SS_AC + SS_BC + SS_ABC)
        
        # DFs
        df_A = self.a - 1
        df_B = self.b - 1
        df_C = self.c - 1
        
        df_AB = df_A * df_B
        df_AC = df_A * df_C
        df_BC = df_B * df_C
        
        df_ABC = df_A * df_B * df_C
        
        df_Error = (self.n - 1) - (df_A + df_B + df_C + df_AB + df_AC + df_BC + df_ABC)
        df_Total = self.n - 1
        
        self.MS_E = SS_Error / df_Error if df_Error > 0 else 0
        self.df_E = df_Error
        
        # Replications
        # r_bar = n / (abc)
        self.r_bar = self.n / (self.a * self.b * self.c)
        
        anova = {}
        
        for name, ss, df in [
            ("Factor A", SS_A, df_A),
            ("Factor B", SS_B, df_B),
            ("Factor C", SS_C, df_C),
            ("Interaction AxB", SS_AB, df_AB),
            ("Interaction AxC", SS_AC, df_AC),
            ("Interaction BxC", SS_BC, df_BC),
            ("Interaction AxBxC", SS_ABC, df_ABC),
            ("Error", SS_Error, df_Error),
        ]:
            if name == "Error":
                anova[name] = {"df": df, "SS": ss, "MS": ss/df if df>0 else 0, "F": None, "P": None}
                continue
                
            ms = ss / df if df > 0 else 0
            f = ms / self.MS_E if self.MS_E > 0 else 0
            p = 1 - stats.f.cdf(f, df, df_Error)
            anova[name] = {"df": df, "SS": ss, "MS": ms, "F": f, "P": p}
            
        anova["Total"] = {"df": df_Total, "SS": SS_Total, "MS": None, "F": None, "P": None}
        self.anova_table = anova
        return anova

    def run_post_hoc(self, method='lsd', alpha=0.05, order='desc'):
        self.alpha = alpha
        results = {}
        is_asc = (order == 'asc')
        
        # Helper for analysis
        def analyze_effect(effect_name, group_cols, divisor_SE_denom, n_means, SE_denom_is_r_bar=False):
            # Means
            if isinstance(group_cols, list) and len(group_cols) > 1:
                # Create phantom column
                col_name = " : ".join(group_cols)
                self.df[col_name] = self.df.apply(lambda x: " : ".join([str(x[c]) for c in group_cols]), axis=1)
                grp_field = col_name
            else:
                grp_field = group_cols[0]
                
            means = self.df.groupby(grp_field)[self.resp_col].mean().sort_values(ascending=is_asc)
            means_display = means.sort_index()
            sds = self.df.groupby(grp_field)[self.resp_col].std().sort_index()
            counts = self.df.groupby(grp_field)[self.resp_col].count().sort_index()
            ses = sds / np.sqrt(counts)
            
            # SE Calculation
            # SEm = sqrt(MS_E / effective_r)
            if SE_denom_is_r_bar:
                 # Interaction or manual r_bar usage
                 # r_eff = divisor_SE_denom (which acts as r count for that mean)
                 # Actually passed arg is just denominator for MS_E
                 # Let's map args strictly to formula:
                 # SEm = sqrt(MS_E / (b*c*r_bar)) for A
                 # Here divisor_SE_denom will be (a * b * r_bar etc)
                 pass
                 
            # Calculating Effective Replication for Mean
            # For Factor A (level i): n_i approx b * c * r_bar
            # We used balanced assumption: r_eff = n_total / n_levels_current
            r_eff = (self.a * self.b * self.c * self.r_bar) / n_means
            
            SEm = np.sqrt(self.MS_E / r_eff)
            SEd = np.sqrt(2) * SEm
            CV = (np.sqrt(self.MS_E) / self.grand_mean) * 100
            
            # Post-hoc Grouping
            grouping =  {k: "ns" for k in means.index}
            if self.anova_table[effect_name]["P"] <= alpha:
                 grouping = self._compute_grouping(means, method, alpha, SEm, self.df_E)
                 
            # CD
            CD = self._get_cd(method, alpha, self.df_E, SEm, n_means)
            
            return {
                "means": means_display, "sds": sds, "ses": ses,
                "grouping": grouping, "SE": SEm, "SEd": SEd, "CV": CV, "CD": CD
            }

        # Main Effects
        results["Factor A"] = analyze_effect("Factor A", [self.a_col], None, self.a)
        results["Factor B"] = analyze_effect("Factor B", [self.b_col], None, self.b)
        results["Factor C"] = analyze_effect("Factor C", [self.c_col], None, self.c)
        
        # Interactions
        results["Interaction AxB"] = analyze_effect("Interaction AxB", [self.a_col, self.b_col], None, self.a * self.b)
        results["Interaction AxC"] = analyze_effect("Interaction AxC", [self.a_col, self.c_col], None, self.a * self.c)
        results["Interaction BxC"] = analyze_effect("Interaction BxC", [self.b_col, self.c_col], None, self.b * self.c)
        results["Interaction AxBxC"] = analyze_effect("Interaction AxBxC", [self.a_col, self.b_col, self.c_col], None, self.a * self.b * self.c)
        
        self.results = results
        return results

    def _get_cd(self, method, alpha, df, SE, n_means=2):
        if method == 'lsd':
            t = stats.t.ppf(1 - alpha/2, df)
            return t * (np.sqrt(2) * SE)
        elif method == 'tukey':
            q = stats.studentized_range.ppf(1-alpha, n_means, df)
            return q * SE
        elif method == 'duncan':
            q = get_duncan_q(2, df, alpha)
            return q * SE
        return None

    def _compute_grouping(self, means, method, alpha, SE, df):
        vals = means.values
        labels = means.index.tolist()
        n = len(vals)
        sig_set = set()
        
        if method == 'lsd':
            limit = stats.t.ppf(1 - alpha/2, df) * np.sqrt(2) * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= limit:
                        sig_set.add((i, j))
        elif method == 'tukey':
            q = stats.studentized_range.ppf(1-alpha, n, df)
            limit = q * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= limit:
                        sig_set.add((i, j))
        elif method == 'duncan':
            for i in range(n):
                for j in range(i+1, n):
                    p = j - i + 1
                    q = get_duncan_q(p, df, alpha)
                    if abs(vals[i] - vals[j]) >= (q * SE):
                        sig_set.add((i, j))
                        
        G = nx.Graph()
        G.add_nodes_from(range(n))
        for i in range(n):
            for j in range(i+1, n):
                if (i, j) not in sig_set:
                    G.add_edge(i, j)
                    
        cliques = list(nx.find_cliques(G))
        cliques.sort(key=lambda c: (-max(c), -len(c)))
        vocab = "abcdefghijklmnopqrstuvwxyz"
        letters = {i: "" for i in range(n)}
        
        for idx, clq in enumerate(cliques):
            if idx < len(vocab):
                char = vocab[idx]
                for node in clq:
                    letters[node] += char
                    
        return {labels[i]: "".join(sorted(letters[i])) for i in range(n)}

    def create_report(self):
        doc = Document()
        doc.add_heading('Three-Factor CRD Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now()}")
        
        # ANOVA
        doc.add_heading('ANOVA Summary', 1)
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = 'Table Grid'
        
        hdr = tbl.rows[0].cells
        for i, h in enumerate(['Source', 'DF', 'SS', 'MS', 'F-val', 'Result']):
            hdr[i].text = h
            
        order = ["Factor A", "Factor B", "Factor C", 
                 "Interaction AxB", "Interaction AxC", "Interaction BxC", 
                 "Interaction AxBxC", "Error", "Total"]
                 
        for k in order:
            if k not in self.anova_table: continue
            row = tbl.add_row().cells
            d = self.anova_table[k]
            row[0].text = k
            row[1].text = str(d['df'])
            row[2].text = f"{d['SS']:.4f}"
            row[3].text = f"{d['MS']:.4f}" if d['MS'] else ""
            row[4].text = f"{d['F']:.4f}" if d['F'] else ""
            if d['P'] is not None:
                sig = "**" if d['P'] <= 0.01 else ("*" if d['P'] <= 0.05 else "ns")
                row[5].text = f"{d['P']:.4f} {sig}"
        
        # Helper for Mean Tables
        def add_res_table(title, key):
            if key not in self.results: return
            doc.add_heading(title, 2)
            res = self.results[key]
            
            doc.add_paragraph(f"SE(m): {res['SE']:.4f} | SE(d): {res['SEd']:.4f} | CV: {res['CV']:.2f}% | CD: {res['CD'] if res['CD'] else 'ns'}")
            
            t = doc.add_table(rows=1, cols=4)
            t.style = 'Table Grid'
            h = t.rows[0].cells
            h[0].text = "Level"
            h[1].text = "Mean"
            h[2].text = "Std err"
            h[3].text = "Group"
            
            for lvl, val in res['means'].items():
                r = t.add_row().cells
                r[0].text = str(lvl)
                r[1].text = f"{val:.4f}"
                r[2].text = f"{res['ses'][lvl]:.4f}"
                r[3].text = res['grouping'][lvl]

        add_res_table("Factor A Means", "Factor A")
        add_res_table("Factor B Means", "Factor B")
        add_res_table("Factor C Means", "Factor C")
        add_res_table("AxB Interaction", "Interaction AxB")
        add_res_table("AxC Interaction", "Interaction AxC")
        add_res_table("BxC Interaction", "Interaction BxC")
        add_res_table("AxBxC Interaction", "Interaction AxBxC")
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
